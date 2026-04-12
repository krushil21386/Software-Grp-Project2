import os
from docx import Document
from docx.shared import Inches, Pt

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
SRS_MD = os.path.join(ROOT, "srs_document.md")
OUT_DIR = os.path.join(ROOT, "docs")
OUT_DOCX = os.path.join(OUT_DIR, "CareSync-SRS.docx")
OUT_DOCX_FINAL = os.path.join(OUT_DIR, "CareSync-SRS-FINAL.docx")

DIAGRAMS = {
    "gantt": os.path.join(ROOT, "docs", "diagrams", "gantt.png"),
    "sequence_booking": os.path.join(ROOT, "docs", "diagrams", "sequence_booking.png"),
    "dfd_l1": os.path.join(ROOT, "docs", "diagrams", "dfd_l1.png"),
    "dfd_l0": os.path.join(ROOT, "docs", "diagrams", "dfd_l0_context.png"),
    "dfd_l2_appt": os.path.join(ROOT, "docs", "diagrams", "dfd_l2_appointment.png"),
}

def add_heading_from_line(doc: Document, line: str):
    if line.startswith("## "):
        doc.add_heading(line.replace("## ", "").strip(), level=1)
        return True
    if line.startswith("### "):
        doc.add_heading(line.replace("### ", "").strip(), level=2)
        return True
    if line.startswith("#### "):
        doc.add_heading(line.replace("#### ", "").strip(), level=3)
        return True
    return False

def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text.rstrip())
    for run in p.runs:
        run.font.size = Pt(11)

def insert_diagram(doc: Document, key: str, title: str):
    path = DIAGRAMS.get(key)
    if path and os.path.exists(path):
        doc.add_paragraph("")
        doc.add_paragraph(title).runs[0].bold = True
        doc.add_picture(path, width=Inches(6.5))
        doc.add_paragraph("")

def main():
    if not os.path.exists(OUT_DIR):
        os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "CareSync SRS"
    doc.core_properties.subject = "Software Requirements Specification"

    in_code_fence = False
    fence_lang = ""

    with open(SRS_MD, "r", encoding="utf-8") as f:
        for raw_line in f:
            line = raw_line.rstrip("\n")

            # Detect mermaid code fences and replace with images
            if line.strip().startswith("```"):
                if not in_code_fence:
                    in_code_fence = True
                    fence_lang = line.strip().strip("`").lower()
                    continue
                else:
                    # closing fence
                    in_code_fence = False
                    fence_lang = ""
                    continue

            if in_code_fence:
                # Skip code body (we render diagrams instead at section markers)
                continue

            # Headings
            if add_heading_from_line(doc, line):
                continue

            # Replace known diagram section titles with images
            lower = line.lower().strip()
            if lower.startswith("### 8.0 context diagram"):
                add_paragraph(doc, " ")
                insert_diagram(doc, "dfd_l0", "Figure 0: Context Diagram (DFD Level 0)")
                continue
            if lower.startswith("### 8.1 timeline"):
                add_paragraph(doc, " ")
                insert_diagram(doc, "gantt", "Figure 1: Project Development Timeline (Gantt)")
                continue
            if lower.startswith("### 8.2 appointment booking"):
                add_paragraph(doc, " ")
                insert_diagram(doc, "sequence_booking", "Figure 2: Appointment Booking Sequence")
                continue
            if lower.startswith("### 8.3 data flow diagram"):
                add_paragraph(doc, " ")
                insert_diagram(doc, "dfd_l1", "Figure 3: Data Flow Diagram (Level 1)")
                continue
            if lower.startswith("### 8.4 dfd level 2"):
                add_paragraph(doc, " ")
                insert_diagram(doc, "dfd_l2_appt", "Figure 4: DFD Level 2 – Appointment Management")
                continue

            # Lists
            if line.strip().startswith("- "):
                doc.add_paragraph(line.strip()[2:], style="List Bullet")
                continue

            # Blank lines
            if line.strip() == "":
                doc.add_paragraph("")
                continue

            # Default paragraph
            add_paragraph(doc, line)

    # Always write FINAL document (fresh filename to avoid locks)
    doc.save(OUT_DOCX_FINAL)
    print(f"Saved FINAL: {OUT_DOCX_FINAL}")

    # Also try to update the main docx (may be locked if open)
    try:
        doc.save(OUT_DOCX)
        print(f"Saved: {OUT_DOCX}")
    except PermissionError:
        import time
        ts = int(time.time())
        alt = OUT_DOCX.replace(".docx", f"-v{ts}.docx")
        doc.save(alt)
        print(f"Target locked. Saved to: {alt}")

if __name__ == "__main__":
    main()

