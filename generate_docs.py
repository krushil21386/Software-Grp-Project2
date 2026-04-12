"""
Medicare Plus — Project Documentation Generator
Generates two DOCX files:
  1. docs/Medicare_Plus_Project_Documentation.docx  (Full Technical Documentation)
  2. docs/Medicare_Plus_Presentation_Guide.docx      (Presentation / Demo Script)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

# ─── STYLE HELPERS ─────────────────────────────────────────────────

CRIMSON   = RGBColor(0xDC, 0x14, 0x3C)
DARK_BLUE = RGBColor(0x1A, 0x73, 0xE8)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY  = RGBColor(0x5F, 0x63, 0x68)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color_hex):
    """Set background colour of a table cell."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = CRIMSON if level <= 2 else DARK_GRAY
    return h

def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run = p.add_run(text)
    else:
        run = p.runs[0] if p.runs else p.add_run(text)
        run.text = text
    run.font.size = Pt(11)
    return p

def add_feature_table(doc, rows):
    """Add a table with Feature | Purpose | Technical Detail columns."""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = 'Feature'
    hdr[1].text = 'Purpose'
    hdr[2].text = 'Technical Implementation'
    for cell in hdr:
        set_cell_shading(cell, 'DC143C')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)

    for feature, purpose, detail in rows:
        row = table.add_row().cells
        row[0].text = feature
        row[1].text = purpose
        row[2].text = detail
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing

def add_api_table(doc, endpoints):
    """Add a table with Method | Endpoint | Auth | Description."""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(['Method', 'Endpoint', 'Auth', 'Description']):
        hdr[i].text = h
        set_cell_shading(hdr[i], '1A73E8')
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)

    for method, endpoint, auth, desc in endpoints:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = auth
        row[3].text = desc
        # Colour code method
        for p in row[0].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                if method == 'POST':
                    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
                elif method == 'GET':
                    run.font.color.rgb = DARK_BLUE
                elif method in ('PUT', 'PATCH'):
                    run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
                elif method == 'DELETE':
                    run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
        for cell in row[1:]:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
#  DOCUMENT 1: FULL PROJECT DOCUMENTATION
# ═══════════════════════════════════════════════════════════════════

def build_full_documentation():
    doc = Document()

    # -- Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # ── TITLE PAGE ──────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Medicare Plus')
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = CRIMSON

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Healthcare Appointment Management System')
    run.font.size = Pt(18)
    run.font.color.rgb = MED_GRAY

    doc.add_paragraph()
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Complete Project Documentation')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    doc.add_paragraph()
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Version 1.0  •  April 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = MED_GRAY

    doc.add_page_break()

    # ── TABLE OF CONTENTS ──────────────────────────────────────────
    add_styled_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Project Overview',
        '2. Technology Stack',
        '3. System Architecture',
        '4. Frontend Documentation',
        '   4.1 Pages & UI Components',
        '   4.2 State Management (Context API)',
        '   4.3 Real-Time Notification System',
        '5. Backend Documentation',
        '   5.1 REST API Endpoints',
        '   5.2 Database Models',
        '   5.3 Service Layer',
        '   5.4 Middleware & Security',
        '6. Key Features — Highlighted',
        '   6.1 AI Medical Report Analysis (Gemini 2.x)',
        '   6.2 Self-Destructing Health Passport',
        '   6.3 Intelligent Appointment Heatmaps',
        '   6.4 Enterprise Security Architecture',
        '   6.5 Automated Email Notification System',
        '7. Deployment & DevOps',
        '8. API Documentation (Swagger)',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ═══════════════════════════════════════════════════════════════
    add_styled_heading(doc, '1. Project Overview', 1)
    add_body(doc, 'Medicare Plus is a full-stack, AI-powered Healthcare Appointment Management System that digitizes the entire patient-doctor lifecycle. It goes beyond a basic booking system by integrating multi-modal AI diagnostics, real-time communication, and a comprehensive security architecture.')
    add_body(doc, 'The platform serves three distinct user roles — Patient, Doctor, and Admin — each with a tailored dashboard experience.')

    add_styled_heading(doc, 'Core Objectives', 2)
    objectives = [
        ('Patient Onboarding: ', 'Secure registration with OTP email verification and profile management.'),
        ('Appointment Management: ', 'Full lifecycle — booking, rescheduling, cancellation — with collision detection and automated email reminders.'),
        ('AI-Powered Diagnostics: ', 'Upload a medical report image and get instant disease detection, lab value extraction, and medicine suggestions powered by Google Gemini 2.x.'),
        ('Health Passport: ', 'A portable, shareable, and downloadable health identity document with QR verification.'),
        ('Real-Time Updates: ', 'WebSocket-based live notifications for appointment status changes and system activity.'),
        ('Admin Oversight: ', 'Comprehensive audit logging, user/doctor management, and system health monitoring.'),
    ]
    for prefix, text in objectives:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 2. TECHNOLOGY STACK
    # ═══════════════════════════════════════════════════════════════
    add_styled_heading(doc, '2. Technology Stack', 1)

    add_styled_heading(doc, 'Frontend', 2)
    fe_stack = [
        ('React 19', 'Core UI library for building component-based SPAs.'),
        ('Vite 7', 'Ultra-fast build tooling with Hot Module Replacement (HMR).'),
        ('React Router v6', 'Client-side routing with protected routes and role-based navigation.'),
        ('Socket.io Client', 'Real-time bidirectional communication with the backend.'),
        ('Recharts', 'Data visualization for doctor performance analytics and heatmaps.'),
        ('Leaflet + React-Leaflet', 'Interactive map integration for hospital/doctor geolocation.'),
        ('jsPDF + jsPDF-AutoTable', 'Client-side PDF generation for the Health Passport feature.'),
        ('QRCode', 'Generates scannable QR codes for document verification.'),
        ('CSS Modules', 'Scoped component styling for maintainable, collision-free CSS.'),
    ]
    add_feature_table(doc, [(t, d, '') for t, d in fe_stack])

    add_styled_heading(doc, 'Backend', 2)
    be_stack = [
        ('Node.js + Express 5', 'High-performance REST API server.'),
        ('MongoDB + Mongoose 9', 'NoSQL database with schema validation and encryption hooks.'),
        ('Socket.io', 'Real-time event broadcasting for live notifications.'),
        ('Google Generative AI (Gemini)', 'Multi-modal AI for medical report image analysis (OCR + LLM).'),
        ('Nodemailer', 'Transactional email service (SMTP/Gmail) for confirmations, reminders, and alerts.'),
        ('JSON Web Tokens (JWT)', 'Stateless authentication with access + refresh token architecture.'),
        ('Helmet', 'HTTP header security hardening (CSP, HSTS, etc.).'),
        ('Multer', 'Secure multipart file upload handling with size/type validation.'),
        ('node-cron', 'Scheduled background jobs for automated email reminders.'),
        ('Redis', 'Performance caching layer with graceful fallback to in-memory mode.'),
        ('Winston', 'Structured, production-grade logging with file and console transports.'),
        ('Swagger (OpenAPI 3.0)', 'Interactive API documentation at /api-docs.'),
        ('bcryptjs', 'Industry-standard password hashing (10+ salt rounds).'),
    ]
    add_feature_table(doc, [(t, d, '') for t, d in be_stack])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 3. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════
    add_styled_heading(doc, '3. System Architecture', 1)
    add_body(doc, 'Medicare Plus follows a modern three-tier architecture with clear separation of concerns:')

    add_styled_heading(doc, 'High-Level Architecture', 2)
    arch_layers = [
        ('Presentation Layer (Frontend): ', 'React SPA served by Vite. Communicates with the backend exclusively via REST API calls and WebSocket events.'),
        ('Application Layer (Backend): ', 'Node.js/Express REST API. Handles business logic, authentication, AI integration, email dispatch, and real-time event broadcasting via Socket.io.'),
        ('Data Layer: ', 'MongoDB for persistent storage (users, appointments, records, audit logs). Redis for caching with automatic fallback to mock memory mode.'),
        ('External Services: ', 'Google Gemini API for AI analysis, Gmail SMTP for transactional emails.'),
    ]
    for prefix, text in arch_layers:
        add_bullet(doc, text, bold_prefix=prefix)

    add_styled_heading(doc, 'Data Flow Summary', 2)
    add_body(doc, 'Patient/Doctor → React Frontend → REST API (Express) → MongoDB')
    add_body(doc, 'Medical Report Upload → Multer → Security Gateway (Malware Scan) → Gemini AI → Structured JSON → Encrypted Storage')
    add_body(doc, 'Appointment Booking → Collision Detection → DB Write → Email Dispatch → Socket.io Broadcast')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 4. FRONTEND DOCUMENTATION
    # ═══════════════════════════════════════════════════════════════
    add_styled_heading(doc, '4. Frontend Documentation', 1)
    add_body(doc, 'The frontend is a React Single-Page Application (SPA) built with Vite. It uses CSS Modules for scoped styling and the Context API for global state management.')

    add_styled_heading(doc, '4.1 Pages & UI Components', 2)
    pages = [
        ('Home.jsx', 'Landing page with hero section, statistics, and top specialists showcase.'),
        ('Login.jsx / Signup.jsx', 'Authentication forms with real-time validation; Login supports role selection (Patient/Doctor/Admin).'),
        ('VerifyOtp.jsx', 'OTP verification screen with resend functionality and countdown timer.'),
        ('PatientDashboard.jsx', 'Patient hub: overview stats, upcoming appointments, quick actions (book, records, AI analysis).'),
        ('DoctorDashboard.jsx', 'Doctor hub: incoming appointments with accept/reject/complete actions, patient records, reschedule modal, real-time analytics charts.'),
        ('AdminDashboard.jsx', 'Admin control center: system stats, user/doctor management tables, system monitoring, and audit log viewer with CSV export.'),
        ('BookAppointmentPage.jsx', 'Multi-step booking flow: search doctors → select slot → confirm details → receive email.'),
        ('Appointments.jsx', 'View all appointments sorted by status (upcoming, completed, cancelled).'),
        ('MedicineAI.jsx', 'AI medicine suggestion engine: upload report or enter symptoms → get diagnosis, lab results table, bounding box overlays, and medicine recommendations.'),
        ('MedicalRecords.jsx', 'Patient: view AI analysis history and prescription details. Doctor: review patient records, add comments, upload/edit prescriptions.'),
        ('HealthPassport.jsx', 'Generates a portable health identity: stats overview, diagnosis history, appointment timeline, prescriptions. Supports PDF download with QR code and 24h shareable link.'),
        ('DoctorLocator.jsx', 'Interactive Leaflet map to find nearby doctors and hospitals with specialization filters.'),
        ('DoctorProfile.jsx', 'Detailed doctor profile with ratings, availability, and booking CTA.'),
        ('DoctorAvailability.jsx', 'Heatmap visualization showing doctor busyness across the week with AI-suggested best booking times.'),
        ('Hospitals.jsx', 'Hospital directory with search and filter capabilities.'),
        ('MedicineDelivery.jsx', 'Browse medicines and place orders with address and payment handling.'),
    ]
    add_feature_table(doc, [(p, d, 'React + CSS Modules') for p, d in pages])

    add_styled_heading(doc, '4.2 State Management (Context API)', 2)
    add_body(doc, 'The application uses two React Context providers for global state:', bold=True)

    add_styled_heading(doc, 'AuthContext', 3)
    auth_features = [
        ('Session Persistence: ', 'Saves user data, access token, and refresh token in localStorage. Restores session automatically on page reload.'),
        ('Login / Register / Logout: ', 'Wraps all authentication API calls and handles token storage.'),
        ('OTP & MFA Verification: ', 'Provides verifyOtp() and verifyMfa() methods for multi-factor authentication flows.'),
        ('Silent Token Refresh: ', 'refreshAccessToken() silently refreshes expired access tokens using the stored refresh token. Logs out the user if refresh fails.'),
        ('authFetch() Wrapper: ', 'A custom fetch wrapper that automatically attaches the Bearer token to every request and retries once on 401 (Unauthorized) after refreshing the token. This ensures seamless API calls without manual token handling in every component.'),
    ]
    for prefix, text in auth_features:
        add_bullet(doc, text, bold_prefix=prefix)

    add_styled_heading(doc, 'NotificationContext', 3)
    notif_features = [
        ('Real-Time WebSocket Listener: ', 'Connects to the backend Socket.io server on mount and listens for live events.'),
        ('appointment_status_update: ', 'Triggers toast notifications when an appointment is accepted, rejected, or completed by a doctor.'),
        ('analytics_update: ', 'Shows subtle system activity notifications to admin/doctor roles when dashboard data changes.'),
        ('Auto-Dismiss: ', 'Notifications are automatically removed after 6 seconds.'),
    ]
    for prefix, text in notif_features:
        add_bullet(doc, text, bold_prefix=prefix)

    add_styled_heading(doc, 'Reusable Components', 3)
    components = [
        ('Navbar', 'Global navigation bar with role-based menu items and logout.'),
        ('Footer', 'Site footer with quick links.'),
        ('Toast / ToastContainer', 'Animated notification popups driven by NotificationContext.'),
        ('ProtectedRoute', 'HOC that guards routes based on authentication status and user role.'),
        ('Skeleton', 'Loading state placeholder for async data.'),
        ('DoctorPerformanceChart', 'Recharts-based analytics visualizations.'),
        ('ProfileImageUpload', 'Avatar upload component with preview.'),
        ('Hero', 'Landing page hero banner with call-to-action.'),
        ('Statistics', 'Animated counters for platform metrics.'),
        ('TopSpecialists', 'Featured doctor cards carousel.'),
    ]
    add_feature_table(doc, [(c, d, 'React component') for c, d in components])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 5. BACKEND DOCUMENTATION
    # ═══════════════════════════════════════════════════════════════
    add_styled_heading(doc, '5. Backend Documentation', 1)
    add_body(doc, 'The backend is a Node.js/Express REST API with MongoDB as the primary data store. It features a modular architecture with controllers, services, models, middleware, and routes.')

    add_styled_heading(doc, '5.1 REST API Endpoints', 2)

    add_styled_heading(doc, 'Authentication (/api/auth)', 3)
    add_api_table(doc, [
        ('POST', '/api/auth/register', 'Public', 'Register a new user (patient/doctor/admin). Sends OTP email.'),
        ('POST', '/api/auth/verify-otp', 'Public', 'Verify OTP for email confirmation or password reset.'),
        ('POST', '/api/auth/resend-otp', 'Public', 'Resend OTP to the user\'s email.'),
        ('POST', '/api/auth/login', 'Public', 'Login with email + password. Returns JWT access + refresh tokens.'),
        ('POST', '/api/auth/verify-mfa', 'Public', 'Verify MFA OTP for suspicious login attempts.'),
        ('POST', '/api/auth/refresh-token', 'Public', 'Exchange refresh token for a new access token.'),
        ('POST', '/api/auth/logout', 'Auth', 'Invalidate refresh token and end session.'),
        ('POST', '/api/auth/forgot-password', 'Public', 'Send password reset OTP (anti-enumeration response).'),
        ('POST', '/api/auth/reset-password', 'Public', 'Reset password with valid OTP.'),
        ('GET', '/api/auth/profile', 'Auth', 'Get current user profile.'),
        ('PUT', '/api/auth/profile', 'Auth', 'Update user profile (name, phone, age, etc.).'),
        ('GET', '/api/auth/doctors', 'Public', 'List all registered doctors.'),
        ('GET', '/api/auth/sessions', 'Auth', 'List active sessions for the logged-in user.'),
    ])

    add_styled_heading(doc, 'Appointments (/api/appointments)', 3)
    add_api_table(doc, [
        ('POST', '/api/appointments/book-appointment', 'Auth', 'Book a new appointment with collision detection.'),
        ('GET', '/api/appointments/my-appointments', 'Auth', 'Get all appointments (grouped by status) for patient or doctor.'),
        ('GET', '/api/appointments/heatmap', 'Auth', 'Get 7×24 appointment density matrix with best-time suggestions.'),
        ('PUT', '/api/appointments/:id/accept', 'Auth', 'Accept a pending appointment (doctor).'),
        ('PUT', '/api/appointments/:id/reject', 'Auth', 'Reject/cancel an appointment (requires reason).'),
        ('PUT', '/api/appointments/:id/complete', 'Auth', 'Mark appointment as completed.'),
        ('PUT', '/api/appointments/:id/reschedule', 'Auth', 'Reschedule with new date/time/reason; re-runs collision check.'),
    ])

    add_styled_heading(doc, 'AI Analysis (/api/ai)', 3)
    add_api_table(doc, [
        ('POST', '/api/ai/analyze', 'Auth', 'Upload medical report image. Returns AI diagnosis, lab values, confidence score, and medicine suggestions.'),
    ])

    add_styled_heading(doc, 'Medical Records (/api/medical-records)', 3)
    add_api_table(doc, [
        ('GET', '/api/medical-records/my-records', 'Patient', 'Get all medical records for the logged-in patient.'),
        ('GET', '/api/medical-records/doctor-records', 'Doctor', 'Get all patient records the doctor has access to.'),
        ('GET', '/api/medical-records/patient/:id', 'Doctor/Admin', 'Get records for a specific patient.'),
        ('PATCH', '/api/medical-records/:id', 'Doctor/Admin', 'Update record status or add doctor comments.'),
        ('PATCH', '/api/medical-records/:id/prescription', 'Doctor', 'Upload a prescription file (multipart).'),
        ('PATCH', '/api/medical-records/:id/prescription-details', 'Doctor', 'Update prescription text details (medicines, dosage, frequency).'),
    ])

    add_styled_heading(doc, 'Health Passport (/api/passport)', 3)
    add_api_table(doc, [
        ('GET', '/api/passport/data', 'Auth', 'Aggregate all health data for passport generation.'),
        ('POST', '/api/passport/share', 'Auth', 'Generate a 24-hour self-destructing share link.'),
        ('GET', '/api/passport/view/:token', 'Public', 'View a shared health passport (read-only, time-limited).'),
    ])

    add_styled_heading(doc, 'Analytics & Admin', 3)
    add_api_table(doc, [
        ('GET', '/api/analytics/stats', 'Auth', 'Get global statistics (total users, appointments, etc.).'),
        ('GET', '/api/analytics/doctor/:name', 'Auth', 'Get performance analytics for a specific doctor.'),
        ('GET', '/api/admin/stats', 'Admin', 'Get admin-level system statistics.'),
        ('GET', '/api/admin/users', 'Admin', 'List all users with role/status filters.'),
        ('GET', '/api/audit/logs', 'Admin', 'Get paginated audit logs.'),
        ('GET', '/api/audit/export', 'Admin', 'Download audit logs as CSV.'),
    ])

    doc.add_page_break()

    # 5.2 Database Models
    add_styled_heading(doc, '5.2 Database Models (MongoDB/Mongoose)', 2)
    models = [
        ('User', 'name, email, password (bcrypt hashed), role (patient/doctor/admin), phone (AES-256 encrypted), address (AES-256 encrypted), age, gender, isVerified, failedLoginAttempts, isLocked, lockUntil, lastLogin, refreshToken, knownLocations, specialty, license, profileImage'),
        ('Appointment', 'userId, doctorId, appointmentId (unique UUID), doctor/patient details, date, time, mode (Online/Offline), status (upcoming/completed/cancelled), isUrgent, rescheduleReason, consultationFee, paymentStatus, reminder24hSent, reminder2hSent'),
        ('MedicalRecord', 'patient (ref: User), reportType, fileName, fileUrl, analysis (AES-256 encrypted JSON), status (pending/reviewed/archived), doctorComments, prescriptionUrl, prescriptionDetails'),
        ('AuditLog', 'userId (ref: User), action, category (AUTH/APPOINTMENT/USER/ADMIN/SYSTEM), status, ipAddress, userAgent, location, details, timestamp'),
        ('Session', 'userId, refreshToken, ipAddress, userAgent, isActive, expiresAt'),
        ('OtpToken', 'email, otp (hashed), type (registration/password-reset/login-mfa), expiresAt'),
        ('Hospital', 'name, address, lat, lng, phone, specialties, rating, traffic'),
        ('Medicine', 'name, category, price, description, stock, manufacturer'),
        ('Order', 'userId, medicines, totalAmount, address, status, paymentMethod'),
        ('Prescription', 'patientId, doctorId, medicines, doctorComments, status, issuedDate'),
        ('Availability', 'doctorId, dayOfWeek, startTime, endTime, isAvailable'),
        ('Department', 'name, description'),
    ]
    add_feature_table(doc, [(m, '', f) for m, f in models])

    # 5.3 Service Layer
    add_styled_heading(doc, '5.3 Service Layer', 2)
    services = [
        ('geminiService.js', 'Core AI engine. Sends medical report images to Google Gemini 2.0 Flash (with fallback to 1.5 Flash) using a detailed multi-phase clinical prompt. Returns structured JSON with confidence score, extracted lab values, disease detection, and medicine suggestions.', 'Uses model cascade pattern: tries 5 Gemini models sequentially. Temperature set to 0.1 for deterministic results. Includes JSON sanitization for malformed AI responses.'),
        ('emailService.js', 'Transactional email service using Nodemailer with Gmail SMTP. Generates rich HTML emails with Google Material Design aesthetics.', 'Supports 7 email types: Appointment Confirmation, Cancellation, Reschedule, 24h Reminder, 2h Reminder, Security Alert (new location login), Prescription Update, and OTP/MFA verification.'),
        ('cronService.js', 'Background job scheduler using node-cron. Runs every 15 minutes to check for upcoming appointments requiring reminder emails.', 'Uses idempotent flags (reminder24hSent, reminder2hSent) to prevent duplicate emails. Parses AM/PM time formats for accurate window calculations.'),
        ('securityGateway.js', 'Multi-purpose security service providing AES-256-CBC encryption/decryption and real-time file malware scanning.', 'Malware scan uses signature-based pattern matching (EICAR, shell_exec, eval) and Shannon entropy analysis to detect packed executables.'),
        ('authService.js', 'JWT token generation (access + refresh) and bcrypt password hashing. Provides user sanitization to strip sensitive fields before API responses.', 'Access tokens expire in 15 minutes; refresh tokens in 7 days.'),
        ('otpService.js', 'OTP generation, email delivery, and verification. Supports multiple OTP types: registration, password-reset, and login-mfa.', 'OTPs are hashed before storage for security. Auto-expires after 10 minutes.'),
        ('securityService.js', 'Location-based risk detection. Checks if a login attempt originates from a recognized location.', 'Uses geoip-lite for IP-to-location resolution. Updates known locations after successful MFA.'),
        ('loggingService.js', 'Centralized audit logging with IP address extraction (request-ip) and geolocation tagging.', 'Non-blocking log writes ensure zero API latency impact.'),
        ('cacheService.js', 'Redis-backed caching layer with automatic fallback to in-memory mode if Redis is unavailable.', 'TTL-based cache invalidation. Prevents redundant database queries for frequently accessed data.'),
        ('prioritizationService.js', 'Risk-based appointment slot prioritization. Detects high-risk conditions from symptom keywords and prioritizes earliest available slots for urgent patients.', 'Uses a keyword dictionary (chest pain, stroke, bleeding, etc.) for clinical risk detection.'),
    ]
    add_feature_table(doc, services)

    doc.add_page_break()

    # 5.4 Middleware
    add_styled_heading(doc, '5.4 Middleware & Security', 2)
    middleware = [
        ('authenticate.js', 'JWT bearer token verification. Extracts user payload from the token and attaches it to req.user for all downstream handlers.', 'Returns 401 on missing/invalid tokens with descriptive error messages.'),
        ('authorize.js', 'Role-Based Access Control (RBAC). Restricts endpoint access to specific roles (e.g., only "doctor" or "admin" can access patient records).', 'Uses a flexible role array: authorize("doctor", "admin").'),
        ('adminMiddleware.js', 'Shorthand middleware for admin-only routes.', 'Checks req.user.role === "admin".'),
        ('rateLimiter.js', 'Express rate limiter for brute-force protection on sensitive auth endpoints.', 'Configurable window and max attempts per IP.'),
        ('optionalAuthenticate.js', 'Allows both authenticated and unauthenticated access. Attaches user if token is present, proceeds without error if not.', 'Used for endpoints like public hospital listings or guest-accessible features.'),
    ]
    add_feature_table(doc, middleware)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 6. KEY FEATURES — HIGHLIGHTED
    # ═══════════════════════════════════════════════════════════════
    add_styled_heading(doc, '6. Key Features — Highlighted', 1)
    add_body(doc, 'This section highlights the unique, differentiating features of Medicare Plus alongside their purpose and technical implementation.', italic=True, color=MED_GRAY)

    # 6.1
    add_styled_heading(doc, '6.1 🧬 AI Medical Report Analysis (Gemini 2.x)', 2)
    add_body(doc, 'PURPOSE: Enable patients and doctors to get instant, AI-powered insights from medical report images without manual data entry.', bold=True, color=CRIMSON)
    ai_details = [
        ('Multi-Modal Input: ', 'Accepts medical report images (JPG/PNG) via file upload. The image is sent directly to Google Gemini as base64-encoded inline data.'),
        ('3-Phase Clinical Prompt: ', 'Phase 1 (Data Extraction) extracts every test name, value, unit, and range. Phase 2 (Clinical Correlation) compares values against WHO/Mayo Clinic standards. Phase 3 (Patient Simplification) translates medical jargon into plain English (e.g., "Hyperglycemia" → "Diabetes").'),
        ('Model Cascade Fallback: ', 'Attempts analysis with 5 Gemini models sequentially (2.0 Flash → 1.5 Flash → 1.5 Flash 8B → experimental models). If all fail, falls back to a rule-based lab analysis engine with 9 built-in biomarker ranges.'),
        ('Visual Bounding Boxes: ', 'On the frontend, abnormal values are highlighted on the uploaded image using HTML5 Canvas overlays — red for abnormal, green for normal.'),
        ('Confidence Score: ', 'Returns a 0-100 diagnostic confidence score with a visual meter (green >80%, amber 50-80%, red <50%).'),
        ('Medicine Suggestions: ', 'Automatically maps detected diseases to recommended medicines with dosage, frequency, duration, and precautions.'),
        ('Automatic Record Persistence: ', 'AI analysis results are automatically encrypted (AES-256) and saved to the MedicalRecord collection for future reference.'),
        ('Malware Scanning: ', 'Before AI analysis, every uploaded file passes through the SecurityGateway which performs signature-based pattern matching and Shannon entropy analysis to block malicious files.'),
    ]
    for prefix, text in ai_details:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()

    # 6.2
    add_styled_heading(doc, '6.2 🛡️ Self-Destructing Health Passport', 2)
    add_body(doc, 'PURPOSE: Provide patients with a portable, verified health identity document that can be securely shared with any healthcare provider.', bold=True, color=CRIMSON)
    passport_details = [
        ('Data Aggregation: ', 'The Health Passport aggregates data from multiple collections: User profile, all Appointments (with status timeline), all MedicalRecords (with AI diagnosis history), and all Prescriptions.'),
        ('PDF Generation: ', 'Generates a professional multi-page PDF using jsPDF on the client side. Includes: patient profile section, diagnosis history table with confidence bars, appointment timeline, prescription details, embedded QR code for verification, and branded headers/footers with page numbering.'),
        ('24-Hour Secure Share: ', 'Generates a temporary, token-based URL that self-destructs after 24 hours. The shared view is read-only and hides sensitive PII like email and phone number.'),
        ('QR Code Verification: ', 'Each generated PDF includes a scannable QR code that links back to the passport verification page.'),
    ]
    for prefix, text in passport_details:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()

    # 6.3
    add_styled_heading(doc, '6.3 📅 Intelligent Appointment Heatmaps', 2)
    add_body(doc, 'PURPOSE: Help patients find the best available time to book and help doctors understand their workload distribution.', bold=True, color=CRIMSON)
    heatmap_details = [
        ('7×24 Density Matrix: ', 'The backend aggregates all non-cancelled appointments for a doctor and generates a 7-day × 24-hour matrix showing appointment density per hour slot.'),
        ('Smart Suggestions: ', 'Analyzes the heatmap to find the least-busy time slots (9 AM – 8 PM, excluding 12–2 PM lunch break) and returns specific "Best Time to Book" recommendations for each day of the week.'),
        ('Collision Detection: ', 'Before booking, the system performs an atomic check: if a doctor already has an "upcoming" appointment at the exact same date and time, it returns a 409 Conflict with a user-friendly message. This also runs during rescheduling.'),
        ('Automated Reminders: ', 'A cron job runs every 15 minutes and sends idempotent email reminders at 24 hours and 2 hours before each appointment.'),
    ]
    for prefix, text in heatmap_details:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()

    # 6.4
    add_styled_heading(doc, '6.4 🔐 Enterprise Security Architecture', 2)
    add_body(doc, 'PURPOSE: Protect patient health information (PHI) with defense-in-depth security across every layer.', bold=True, color=CRIMSON)
    security_details = [
        ('AES-256-CBC Encryption at Rest: ', 'Sensitive fields (phone, address) in the User model and the entire AI analysis JSON in MedicalRecord are encrypted using AES-256-CBC before being stored in MongoDB. Decryption happens transparently via Mongoose post-init hooks.'),
        ('bcrypt Password Hashing: ', 'All passwords are hashed with bcrypt (10+ salt rounds) before storage. Plain text is never stored.'),
        ('JWT Access + Refresh Token Architecture: ', 'Short-lived access tokens (15 min) paired with long-lived refresh tokens (7 days). The frontend silently refreshes tokens using the authFetch() wrapper.'),
        ('Account Soft-Locking: ', '5 failed login attempts trigger a 30-minute account lock. The system tracks failedLoginAttempts and lockUntil with automatic unlock after expiry.'),
        ('Risk-Based MFA (Prepared): ', 'The backend is fully wired for location-based MFA: when a login from an unrecognized IP/location is detected, it can trigger an OTP verification step before granting access.'),
        ('Tamper-Evident Audit Logging: ', 'Every critical action (login, booking, record access) is logged with userId, action, category, IP address, user agent, and timestamp. Admins can view and export logs as CSV.'),
        ('File Malware Scanning: ', 'Every uploaded file passes through the SecurityGateway for EICAR signature matching and entropy analysis before being processed.'),
        ('Helmet Security Headers: ', 'The Express app is hardened with Helmet, enforcing Content Security Policy, X-Frame-Options, X-Content-Type-Options, and Referrer-Policy headers.'),
        ('Anti-Enumeration: ', 'The forgot-password endpoint always responds with a generic message regardless of whether the email exists, preventing email enumeration attacks.'),
    ]
    for prefix, text in security_details:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()

    # 6.5
    add_styled_heading(doc, '6.5 📧 Automated Email Notification System', 2)
    add_body(doc, 'PURPOSE: Keep patients informed at every step with professional, branded email communications.', bold=True, color=CRIMSON)
    email_details = [
        ('Appointment Confirmation: ', 'Rich HTML email with appointment ID badge, doctor details, date/time, mode (Online/Offline), patient details, Google Maps deep-link, and branded footer.'),
        ('Cancellation Notification: ', 'Red-themed alert with the specific cancellation reason provided by the doctor.'),
        ('Reschedule Notification: ', 'Side-by-side comparison of old vs. new date/time with reason.'),
        ('24h & 2h Reminders: ', 'Automated via cron job. Uses idempotent flags to prevent duplicate sends. 2h reminder uses an urgent orange theme.'),
        ('Security Alert: ', 'Notifies users when a login from a new, unrecognized location is detected, showing IP and location.'),
        ('Prescription Update: ', 'Notifies patients when a doctor uploads or updates their prescription, with a direct link to view records.'),
        ('OTP/MFA Codes: ', 'Clean, centered OTP display with dashed border and expiry warning.'),
    ]
    for prefix, text in email_details:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 7. DEPLOYMENT
    # ═══════════════════════════════════════════════════════════════
    add_styled_heading(doc, '7. Deployment & DevOps', 1)

    add_styled_heading(doc, 'Local Development', 2)
    add_body(doc, 'Backend:')
    add_bullet(doc, 'cd backend && npm install && npm run dev')
    add_body(doc, 'Frontend:')
    add_bullet(doc, 'cd Software-Grp-Project && npm install && npm run dev')

    add_styled_heading(doc, 'Environment Variables (backend/.env)', 2)
    env_vars = [
        ('PORT', '5000'),
        ('MONGO_URI', 'mongodb://localhost:27017/healthcare'),
        ('JWT_SECRET', 'Your JWT signing secret'),
        ('JWT_REFRESH_SECRET', 'Your refresh token secret'),
        ('EMAIL_USER', 'Gmail address for SMTP'),
        ('EMAIL_PASS', 'Gmail App Password'),
        ('GEMINI_API_KEY', 'Google AI Studio API key'),
        ('ENCRYPTION_KEY', '32-byte AES encryption key'),
        ('REDIS_URL', 'redis://localhost:6379 (optional)'),
    ]
    add_feature_table(doc, [(v, d, '') for v, d in env_vars])

    add_styled_heading(doc, 'Docker Deployment', 2)
    add_body(doc, 'The project includes a docker-compose.yml for containerized deployment:')
    add_bullet(doc, 'docker-compose up --build')
    add_bullet(doc, 'Frontend: http://localhost:5173')
    add_bullet(doc, 'Backend: http://localhost:5000')

    add_styled_heading(doc, 'Seed Data', 2)
    add_bullet(doc, 'node backend/seedAdmin.js — Creates a default admin account')
    add_bullet(doc, 'node backend/seedDoctors.js — Populates sample doctor profiles')
    add_bullet(doc, 'node backend/seedData.js — Seeds hospitals, departments, and test data')
    add_bullet(doc, 'node backend/seedMedicines.js — Populates the medicine catalog')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 8. API DOCUMENTATION
    # ═══════════════════════════════════════════════════════════════
    add_styled_heading(doc, '8. API Documentation (Swagger)', 1)
    add_body(doc, 'Medicare Plus includes a built-in interactive API documentation powered by Swagger UI (OpenAPI 3.0).')
    add_body(doc, 'Access it at: http://localhost:5000/api-docs', bold=True, color=DARK_BLUE)
    add_body(doc, 'The Swagger UI provides:')
    swagger_features = [
        'Try-it-out functionality for every endpoint',
        'JWT Bearer authentication support',
        'Request/response schema documentation',
        'Organized by tags: Auth, AI Analysis, Appointments, Medical Records, Health Passport, Analytics, Audit',
        'File upload testing (multipart/form-data)',
    ]
    for f in swagger_features:
        add_bullet(doc, f)

    # ── SAVE ───────────────────────────────────────────────────────
    os.makedirs('docs', exist_ok=True)
    path = os.path.join('docs', 'Medicare_Plus_Project_Documentation.docx')
    doc.save(path)
    print(f'✅ Saved: {path}')
    return path


# ═══════════════════════════════════════════════════════════════════
#  DOCUMENT 2: PRESENTATION GUIDE
# ═══════════════════════════════════════════════════════════════════

def build_presentation_guide():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # ── TITLE PAGE ──────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Medicare Plus')
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = CRIMSON

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Project Presentation Guide')
    run.font.size = Pt(18)
    run.font.color.rgb = MED_GRAY

    doc.add_paragraph()
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Demo Script  •  Talking Points  •  Q&A Preparation')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Version 1.0  •  April 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = MED_GRAY

    doc.add_page_break()

    # ── SLIDE 1: INTRODUCTION ─────────────────────────────────────
    add_styled_heading(doc, 'Slide 1: Project Introduction', 1)
    add_body(doc, 'Title: "Medicare Plus — AI-Powered Healthcare Management"', bold=True)
    add_styled_heading(doc, 'Key Talking Points', 3)
    intro_points = [
        '"Medicare Plus is a full-stack Healthcare Appointment Management System that we built from the ground up."',
        '"It goes beyond a basic booking app — it integrates AI-powered diagnostics, real-time notifications, and enterprise-grade security."',
        '"The system serves 3 user roles: Patient, Doctor, and Admin — each with a tailored experience."',
    ]
    for p in intro_points:
        add_bullet(doc, p)

    add_styled_heading(doc, 'Tech Stack to Mention', 3)
    add_body(doc, 'Frontend: React 19 + Vite 7 | Backend: Node.js + Express 5 | Database: MongoDB | AI: Google Gemini 2.x | Real-time: Socket.io | Email: Nodemailer')

    doc.add_paragraph()

    # ── SLIDE 2: ARCHITECTURE ─────────────────────────────────────
    add_styled_heading(doc, 'Slide 2: System Architecture', 1)
    add_body(doc, 'Title: "Three-Tier Architecture with AI Integration"', bold=True)
    add_styled_heading(doc, 'Key Talking Points', 3)
    arch_points = [
        '"The frontend is a React SPA that communicates with a RESTful Express backend."',
        '"MongoDB handles all persistent data with encrypted fields for PHI compliance."',
        '"Real-time updates flow through Socket.io WebSockets — no polling required."',
        '"We integrated Google Gemini 2.0 Flash for multi-modal medical report analysis."',
        '"The system uses Redis for caching, with automatic fallback to in-memory mode."',
    ]
    for p in arch_points:
        add_bullet(doc, p)

    add_styled_heading(doc, 'Diagram to Show', 3)
    add_body(doc, 'Show the DFD Level 0: Patient/Doctor/Admin → REST API + Socket.io → MongoDB + Gemini AI', italic=True, color=MED_GRAY)

    doc.add_paragraph()

    # ── SLIDE 3: DEMO - AUTH ──────────────────────────────────────
    add_styled_heading(doc, 'Slide 3: Live Demo — Authentication Flow', 1)
    add_body(doc, 'Title: "Secure Onboarding with OTP Verification"', bold=True)
    add_styled_heading(doc, 'Demo Steps', 3)
    auth_demo = [
        '1. Open the Signup page → Fill in the registration form as a Patient.',
        '2. Submit → Show the "OTP sent" success message.',
        '3. Open your email inbox → Show the branded OTP email.',
        '4. Enter the OTP on the verification page → Account verified.',
        '5. Login with the new account → Show the JWT-based session creation.',
        '6. (Optional) Demonstrate failed login attempts and account locking after 5 failures.',
    ]
    for step in auth_demo:
        add_bullet(doc, step)

    add_styled_heading(doc, 'Key Points to Highlight', 3)
    auth_highlights = [
        '"OTP is sent via Gmail SMTP with a professional HTML email template."',
        '"Passwords are hashed with bcrypt — never stored in plain text."',
        '"We implement JWT access + refresh tokens with silent token refresh on the frontend."',
        '"After 5 failed login attempts, the account is soft-locked for 30 minutes."',
    ]
    for h in auth_highlights:
        add_bullet(doc, h)

    doc.add_paragraph()

    # ── SLIDE 4: DEMO - BOOKING ───────────────────────────────────
    add_styled_heading(doc, 'Slide 4: Live Demo — Appointment Booking', 1)
    add_body(doc, 'Title: "Smart Booking with Collision Detection & Heatmaps"', bold=True)
    add_styled_heading(doc, 'Demo Steps', 3)
    booking_demo = [
        '1. Login as a Patient → Navigate to "Book Appointment".',
        '2. Search for a doctor by specialization → Select a doctor.',
        '3. Pick a date and time → Click "Book".',
        '4. Show the confirmation toast and the confirmation email in your inbox.',
        '5. Try booking the SAME doctor at the SAME time → Show the 409 Conflict error.',
        '6. Navigate to "Doctor Availability" → Show the 7×24 heatmap visualization.',
        '7. Point out the "Best Time Suggestions" panel.',
    ]
    for step in booking_demo:
        add_bullet(doc, step)

    add_styled_heading(doc, 'Key Points to Highlight', 3)
    booking_highlights = [
        '"We perform atomic collision detection — no two patients can book the same slot."',
        '"The heatmap is a real 7×24 matrix built from actual appointment data."',
        '"Automated email reminders are sent at 24h and 2h before appointments via cron jobs."',
        '"All booking actions emit Socket.io events for real-time dashboard updates."',
    ]
    for h in booking_highlights:
        add_bullet(doc, h)

    doc.add_paragraph()

    # ── SLIDE 5: DEMO - AI ────────────────────────────────────────
    add_styled_heading(doc, 'Slide 5: Live Demo — AI Medical Report Analysis ⭐', 1)
    add_body(doc, 'Title: "Gemini 2.x Multi-Modal Diagnostics"', bold=True)
    add_body(doc, '⭐ THIS IS THE WOW FACTOR DEMO — Spend extra time here.', bold=True, color=CRIMSON)
    add_styled_heading(doc, 'Demo Steps', 3)
    ai_demo = [
        '1. Login as a Patient → Navigate to "AI Medicine Suggestion".',
        '2. Upload a sample medical report image (blood test / lab report).',
        '3. Click "Get Medicine Suggestions" → Wait for Gemini analysis (5-15 seconds).',
        '4. Show the results: Primary Diagnosis with confidence score.',
        '5. Scroll to the Lab Results table — show extracted values with Normal/High/Low status.',
        '6. Point out the bounding boxes on the uploaded image (green = normal, red = abnormal).',
        '7. Show the Recommended Medicines section with dosage, frequency, and precautions.',
        '8. Expand "Show Raw Extracted Text" to demonstrate OCR accuracy.',
    ]
    for step in ai_demo:
        add_bullet(doc, step)

    add_styled_heading(doc, 'Key Points to Highlight', 3)
    ai_highlights = [
        '"This uses Google Gemini 2.0 Flash — the latest multi-modal AI model."',
        '"We send the raw image to Gemini with a 3-phase clinical prompt: Extract → Correlate → Simplify."',
        '"The system maps medical jargon to plain English — \'Hyperglycemia\' becomes \'Diabetes\'."',
        '"If Gemini is down, it automatically falls back to our rule-based engine with 9 biomarker ranges."',
        '"Every analysis result is encrypted with AES-256 before being saved to the database."',
        '"Uploaded files are scanned for malware before processing."',
    ]
    for h in ai_highlights:
        add_bullet(doc, h)

    doc.add_paragraph()

    # ── SLIDE 6: DEMO - HEALTH PASSPORT ───────────────────────────
    add_styled_heading(doc, 'Slide 6: Live Demo — Health Passport', 1)
    add_body(doc, 'Title: "Portable Health Identity with QR Verification"', bold=True)
    add_styled_heading(doc, 'Demo Steps', 3)
    passport_demo = [
        '1. Navigate to "Health Passport" → Show the aggregated health summary.',
        '2. Click "Download Health Passport PDF" → Open the generated PDF.',
        '3. Show the PDF contents: profile, diagnosis history, appointments, prescriptions, QR code.',
        '4. Click "Generate 24h Share Link" → Copy the link.',
        '5. Open the link in an incognito window → Show the read-only shared view.',
        '6. Point out that email/phone are hidden in the shared view for privacy.',
    ]
    for step in passport_demo:
        add_bullet(doc, step)

    add_styled_heading(doc, 'Key Points to Highlight', 3)
    passport_highlights = [
        '"This is generated entirely on the client side using jsPDF — no server load."',
        '"The share link self-destructs after 24 hours for security."',
        '"The PDF includes an embedded QR code for document verification."',
    ]
    for h in passport_highlights:
        add_bullet(doc, h)

    doc.add_paragraph()

    # ── SLIDE 7: DEMO - DOCTOR DASHBOARD ──────────────────────────
    add_styled_heading(doc, 'Slide 7: Live Demo — Doctor Dashboard', 1)
    add_body(doc, 'Title: "Doctor Workflow: Review, Prescribe, Analyze"', bold=True)
    add_styled_heading(doc, 'Demo Steps', 3)
    doctor_demo = [
        '1. Login as a Doctor → Show the Doctor Dashboard.',
        '2. View incoming appointments → Accept one and reject one (with reason).',
        '3. Show the real-time notification that appears on the patient\'s browser.',
        '4. Open a patient\'s medical record → Show the AI analysis from earlier.',
        '5. Add doctor comments and upload/edit a prescription.',
        '6. Show the prescription notification email sent to the patient.',
        '7. View the doctor\'s performance analytics chart.',
    ]
    for step in doctor_demo:
        add_bullet(doc, step)

    doc.add_paragraph()

    # ── SLIDE 8: DEMO - ADMIN ─────────────────────────────────────
    add_styled_heading(doc, 'Slide 8: Live Demo — Admin Dashboard', 1)
    add_body(doc, 'Title: "Admin Control Center & Audit Trail"', bold=True)
    add_styled_heading(doc, 'Demo Steps', 3)
    admin_demo = [
        '1. Login as Admin → Show the system overview with real stats from the database.',
        '2. Navigate to "Manage Users" → Show the user list with role badges.',
        '3. Navigate to "Manage Doctors" → Show doctor cards.',
        '4. Navigate to "Audit Logs" → Show the action log with timestamps, IPs, and user agents.',
        '5. Click "Download Full CSV" → Show the exported audit log file.',
    ]
    for step in admin_demo:
        add_bullet(doc, step)

    doc.add_paragraph()

    # ── SLIDE 9: SECURITY ─────────────────────────────────────────
    add_styled_heading(doc, 'Slide 9: Security Architecture', 1)
    add_body(doc, 'Title: "Defense in Depth — Every Layer Protected"', bold=True)
    add_styled_heading(doc, 'Key Points', 3)
    security_points = [
        '"PHI (phone, address) is encrypted at rest using AES-256-CBC with automatic Mongoose hooks."',
        '"AI analysis results are encrypted before database storage."',
        '"JWT access tokens expire in 15 minutes; refresh tokens in 7 days."',
        '"5 failed logins = 30-minute account lock."',
        '"Every uploaded file is scanned for malware signatures and suspicious entropy."',
        '"All actions are audit-logged with IP, user agent, and geolocation."',
        '"HTTP headers are hardened with Helmet (CSP, HSTS, X-Frame-Options)."',
        '"Forgot-password uses anti-enumeration responses."',
    ]
    for p in security_points:
        add_bullet(doc, p)

    doc.add_paragraph()

    # ── SLIDE 10: WRAP UP ─────────────────────────────────────────
    add_styled_heading(doc, 'Slide 10: Summary & Conclusion', 1)
    add_body(doc, 'Title: "Medicare Plus — Not Just a Booking System"', bold=True)
    add_styled_heading(doc, 'Closing Statement', 3)
    add_body(doc, '"Medicare Plus is a comprehensive Medical Intelligence Platform that combines AI-powered diagnostics, real-time communication, and enterprise-grade security into a single, cohesive system. It serves patients, doctors, and administrators with tailored experiences — from booking an appointment to generating an encrypted, shareable Health Passport."')

    add_styled_heading(doc, 'Unique Differentiators to Reiterate', 3)
    differentiators = [
        '✅ Multi-modal AI analysis with Gemini 2.x and rule-based fallback',
        '✅ Self-destructing Health Passport with QR verification',
        '✅ 7×24 appointment heatmaps with smart time suggestions',
        '✅ AES-256 encryption for PHI at rest',
        '✅ 7 types of automated, branded transactional emails',
        '✅ Real-time notifications via WebSockets',
        '✅ Interactive API docs via Swagger UI',
        '✅ Docker-ready deployment',
    ]
    for d in differentiators:
        add_bullet(doc, d)

    doc.add_page_break()

    # ── Q&A PREPARATION ───────────────────────────────────────────
    add_styled_heading(doc, 'Appendix: Anticipated Q&A', 1)
    qa = [
        ('Q: What happens if the Gemini AI is down?',
         'A: The system has a cascading fallback. It tries 5 different Gemini model versions, and if all fail, it switches to a built-in rule-based analysis engine that covers 9 common biomarkers (hemoglobin, glucose, WBC, cholesterol, etc.) with standard medical ranges.'),
        ('Q: How do you handle data security for medical records?',
         'A: We implement defense-in-depth: passwords are bcrypt-hashed, PHI fields (phone, address) are AES-256 encrypted at rest via Mongoose hooks, AI analysis data is encrypted before storage, JWT tokens have short expiry (15 min), and all critical actions are audit-logged with IP and user agent.'),
        ('Q: How does the appointment collision detection work?',
         'A: Before creating a booking, we query for any existing "upcoming" appointment with the same doctorId + date + time. If found, we return a 409 Conflict. The same check runs during rescheduling, excluding the current appointment being rescheduled.'),
        ('Q: How do reminder emails work?',
         'A: A node-cron job runs every 15 minutes. It queries all "upcoming" appointments and checks if the appointment is within the 24h or 2h window. It uses idempotent boolean flags (reminder24hSent, reminder2hSent) to ensure each reminder is sent exactly once.'),
        ('Q: Can the Health Passport be forged?',
         'A: The PDF includes an embedded QR code that links back to the verification page. The share link is cryptographically generated with a 24-hour expiry. Shared views are read-only and hide sensitive PII.'),
        ('Q: Why did you choose Context API over Redux?',
         'A: For the scale of this project, Context API with useCallback for memoization provides clean, sufficient state management without the Redux boilerplate. We have two focused contexts (Auth + Notifications) rather than one large global store.'),
        ('Q: How does the real-time notification system work?',
         'A: The NotificationContext establishes a Socket.io connection on mount. The backend emits events like "appointment_status_update" and "analytics_update". The frontend listens, creates toast notifications, and auto-dismisses them after 6 seconds.'),
        ('Q: What is the purpose of Redis in your project?',
         'A: Redis is used for performance caching of frequently accessed data. However, it\'s designed as an optional dependency — if Redis is unavailable, the system automatically falls back to a mock in-memory client with zero downtime.'),
    ]
    for question, answer in qa:
        add_body(doc, question, bold=True, color=CRIMSON)
        add_body(doc, answer)
        doc.add_paragraph()

    # ── SAVE ───────────────────────────────────────────────────────
    path = os.path.join('docs', 'Medicare_Plus_Presentation_Guide.docx')
    doc.save(path)
    print(f'✅ Saved: {path}')
    return path


# ═══════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    print('📄 Generating Medicare Plus Documentation Suite...\n')
    build_full_documentation()
    build_presentation_guide()
    print('\n🎉 All documents generated successfully in the docs/ folder!')
